from pathlib import Path

import pytest
from docx import Document

from document_enhancer.ingest import SourceIngestionError, ingest_source
from document_enhancer.models import SourceFormat


def test_markdown_ingestion_preserves_complete_order_and_details(tmp_path: Path) -> None:
    source = tmp_path / "procedure.md"
    source.write_text(
        """Before the title, retain this control note.

# Monthly Access Review

## Purpose

The analyst confirms that active accounts remain authorized.

## Actions

1. Export the Active Accounts report from Orion.
2. Compare every account with the approved roster.
3. Save the reviewed report as `access-review.csv`.
""",
        encoding="utf-8",
    )

    result = ingest_source(source)

    assert result.source_format is SourceFormat.MARKDOWN
    assert result.title == "Monthly Access Review"
    assert [block.id for block in result.blocks] == ["SRC-001", "SRC-002", "SRC-003"]
    assert [block.heading for block in result.blocks] == ["Preamble", "Purpose", "Actions"]
    assert "Compare every account" in result.blocks[-1].content
    assert "Save the reviewed report" in result.full_text


def test_docx_ingestion_preserves_lists_and_tables_in_document_order(tmp_path: Path) -> None:
    source = tmp_path / "procedure.docx"
    document = Document()
    document.add_heading("Weekly Exception Review", level=0)
    document.add_heading("Actions", level=1)
    document.add_paragraph("Open the exception queue in Atlas.", style="List Number")
    document.add_paragraph("Filter the queue to the current week.", style="List Number")
    document.add_heading("Evidence", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Record"
    table.cell(0, 1).text = "Location"
    table.cell(1, 0).text = "Review log"
    table.cell(1, 1).text = "Atlas"
    document.save(source)

    result = ingest_source(source)

    assert result.source_format is SourceFormat.DOCX
    assert result.title == "Weekly Exception Review"
    assert [block.heading for block in result.blocks] == ["Actions", "Evidence"]
    assert "1. Open the exception queue" in result.blocks[0].content
    assert "2. Filter the queue" in result.blocks[0].content
    assert "| Record | Location |" in result.blocks[1].content
    assert "| Review log | Atlas |" in result.blocks[1].content


@pytest.mark.parametrize("name", ["source.txt", "source.pdf"])
def test_unsupported_source_format_has_clear_error(tmp_path: Path, name: str) -> None:
    source = tmp_path / name
    source.write_text("content", encoding="utf-8")
    with pytest.raises(SourceIngestionError, match=r"expected \.docx, \.md, or \.markdown"):
        ingest_source(source)


def test_empty_markdown_is_rejected(tmp_path: Path) -> None:
    source = tmp_path / "empty.md"
    source.write_text("\n", encoding="utf-8")
    with pytest.raises(SourceIngestionError, match="empty"):
        ingest_source(source)
