"""Build the realistic DOCX source used by the document-enhancer evaluation tests.

The fixture is intentionally a little untidy: it contains metadata tables, mixed
paragraph styles, an old desk note, and repeated unresolved statements.  It still
describes one complete desktop procedure so the ingestion and analysis stages can
be tested against both usable facts and bounded defects.

Run from any working directory with::

    uv run python tests/fixtures/build_messy_fixture.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches

FIXTURE_DIR = Path(__file__).resolve().parent
DOCX_PATH = FIXTURE_DIR / "messy_desktop_procedure.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    """Replace a table cell's text while keeping a predictable run style."""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a compact, readable table whose contents survive source ingestion."""

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            _set_cell_text(cell, value)


def _add_labelled_paragraph(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)


def _add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def _add_numbered_step(document: Document, text: str) -> None:
    # Keep all top-level actions contiguous.  The ingestion contract can then
    # preserve their source order as 1..9 even when the source is regenerated.
    document.add_paragraph(text, style="List Number")


def build_fixture(output_path: Path = DOCX_PATH) -> Path:
    """Create and return the deterministic messy desktop-procedure DOCX."""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    document.core_properties.title = "Weekly Purchase-Order Invoice Variance Review"
    document.core_properties.subject = "Internal desktop procedure for AP variance review"
    document.core_properties.author = "Finance Operations"
    document.core_properties.comments = "Synthetic evaluation fixture; no live business data."

    title = document.add_paragraph(style="Title")
    title.add_run("Weekly Purchase-Order Invoice Variance Review")
    subtitle = document.add_paragraph()
    subtitle.add_run(
        "Internal work instruction | Finance Operations | synthetic evaluation fixture"
    ).italic = True

    _add_table(
        document,
        ["Field", "Value"],
        [
            ["Document status", "Working instruction"],
            ["Audience", "North America Accounts Payable analysts"],
            ["Cadence", "Every Tuesday after 09:00 ET"],
            ["Source of truth", "LedgerHub export plus approved supplier roster"],
            ["Document owner", "Finance Operations"],
        ],
    )
    document.add_paragraph(
        "Use this checklist for the current week's PO-backed invoice variance file. "
        "The screen labels and folder names below are the ones used by the desktop team."
    )

    document.add_heading("Purpose, audience, and scope", level=1)
    document.add_paragraph(
        "The weekly review identifies unmatched or over-threshold purchase-order invoice "
        "variances before Accounts Payable posts the payment batch. Successful completion "
        "produces a reviewed workbook, exception cases for unresolved rows, and a retained "
        "audit package."
    )
    _add_labelled_paragraph(
        document,
        "Audience: ",
        "Accounts Payable analysts who process North America purchase-order invoices in LedgerHub.",
    )
    _add_labelled_paragraph(
        document,
        "Included: ",
        "The current week's PO-backed invoice export, supplier-roster comparison, variance "
        "classification, exception cases, reviewer validation, and evidence handoff.",
    )
    _add_labelled_paragraph(
        document,
        "Excluded: ",
        "Non-PO invoices, credit memos, and historical weeks unless a Finance Reviewer "
        "reopens them.",
    )

    document.add_heading("Prerequisites, access, and inputs", level=1)
    document.add_paragraph(
        "Before starting, confirm that the export has arrived, the weekly folder is writable, "
        "and LedgerHub is available. Use SSO and MFA; do not store credentials in the workbook."
    )
    _add_table(
        document,
        ["Category", "Required item", "Use in this procedure"],
        [
            ["Access", "LedgerHub", "Open Invoice Match and review invoice records."],
            ["Access", "Procurement shared mailbox", "Download the current PO variance export."],
            [
                "Access",
                "Finance shared drive and Exception Queue",
                "Save evidence and create exception cases.",
            ],
            ["Tools", "Excel desktop and Teams", "Calculate, review, and hand off the package."],
            [
                "Input file",
                "PO_Variance_YYYYMMDD.csv",
                "Current week's source export; keep the original unchanged.",
            ],
            [
                "Input file",
                "approved_supplier_roster.xlsx",
                "Check Supplier IDs before classifying rows.",
            ],
            [
                "Input record",
                "Invoice PDF in LedgerHub",
                "Attach supporting evidence for exception rows.",
            ],
            ["Output template", "review_log.csv", "Record reviewer initials and review timestamp."],
        ],
    )
    _add_bullet(
        document, "Use Finance\\AP\\Variance\\YYYY-WW\\Incoming for the untouched source copy."
    )
    _add_bullet(
        document, "Do not begin if the LedgerHub period is locked; use the exception path instead."
    )

    document.add_heading("Roles, trigger, and timing", level=1)
    _add_table(
        document,
        ["Role", "Responsibility", "Handoff"],
        [
            [
                "Accounts Payable analyst",
                "Downloads, compares, classifies, and packages the review.",
                "Sends the evidence link to the Finance Reviewer.",
            ],
            [
                "Procurement analyst",
                "Publishes the PO_Variance_YYYYMMDD.csv export and corrects a missing export.",
                "Confirms a replacement export in the shared mailbox.",
            ],
            [
                "Finance Reviewer",
                "Checks classifications, case IDs, and validation evidence.",
                "Provides reviewer initials and timestamp in review_log.csv.",
            ],
        ],
    )
    document.add_paragraph(
        "Trigger: start every Tuesday after the Procurement analyst publishes the current export "
        "and repeat the review if a revised export arrives before the deadline. Complete the "
        "standard path and send the evidence-folder link before 16:00 ET the same day."
    )
    document.add_paragraph(
        "The exception paths refer to a 'finance owner,' but this checklist does not identify "
        "the person's name, team queue, or contact method. Do not invent that missing owner."
    )

    document.add_heading("Desktop procedure — perform in order", level=1)
    _add_numbered_step(
        document,
        "At the Tuesday trigger, open the Procurement shared mailbox and download the latest "
        "PO_Variance_YYYYMMDD.csv export. Confirm the email timestamp is after 09:00 ET; copy "
        "the original, without editing cells, to Finance\\AP\\Variance\\YYYY-WW\\Incoming.",
    )
    _add_numbered_step(
        document,
        "Open LedgerHub with SSO and MFA, choose Accounts Payable > Invoice Match, and set the "
        "business unit to North America and period to YYYY-WW. Confirm the period is not locked "
        "before changing any status.",
    )
    _add_numbered_step(
        document,
        "Open approved_supplier_roster.xlsx from the same weekly folder. Compare each Supplier "
        "ID in the CSV with the roster; mark an unmatched ID as MISSING_SUPPLIER instead of "
        "adding it to the roster.",
    )
    _add_numbered_step(
        document,
        "In a copy of the export, add variance_amount = invoice_amount - po_amount for each row. "
        "Save the working file as PO_Variance_YYYYMMDD_reviewed.xlsx; do not overwrite the "
        "original CSV.",
    )
    _add_numbered_step(
        document,
        "Filter the working file for non-zero variance_amount, a blank PO number, or duplicate "
        "invoice number. Apply reason codes AMOUNT, MISSING_PO, or DUPLICATE, and attach the "
        "related invoice PDF from LedgerHub to the weekly review folder.",
    )
    _add_numbered_step(
        document,
        "Read the threshold in the LedgerHub decision panel: an absolute total variance at or "
        "below $25 is eligible for Clear; anything above it goes to Pending review. Do not "
        "silently substitute the $30 value in the old desk note.",
    )
    _add_numbered_step(
        document,
        "For every Pending row, create an exception case in Exception Queue, copy its case ID "
        "into the working file, and assign the case to the Finance Reviewer. If Exception Queue "
        "does not return a case ID, leave the row Pending and use the failure path below.",
    )
    _add_numbered_step(
        document,
        "Ask the Finance Reviewer to inspect the reason code, invoice PDF, and case ID. Record "
        "reviewer initials and the review timestamp in review_log.csv; keep unresolved rows "
        "Pending.",
    )
    _add_numbered_step(
        document,
        "Save the reviewed workbook, review_log.csv, case export, and a screenshot named "
        "variance_summary_YYYYMMDD.png under Finance\\AP\\Variance\\YYYY-WW\\Evidence. Send "
        "the evidence-folder link to the Finance Reviewer before 16:00 ET.",
    )

    document.add_heading("Decision point and validation", level=1)
    document.add_paragraph(
        "Use the conditions below in order. A row that meets more than one condition remains "
        "Pending until the Finance Reviewer confirms the classification."
    )
    _add_table(
        document,
        ["Decision condition", "Required action", "Expected result"],
        [
            [
                "Absolute total variance <= $25 and no other exception",
                "Select Clear.",
                "Status is Clear and the row is included in review_log.csv.",
            ],
            [
                "Absolute total variance > $25, missing PO, duplicate invoice, or unmatched "
                "supplier",
                "Keep Pending, create an Exception Queue case, and assign Finance Reviewer.",
                "Case ID appears in the working file.",
            ],
        ],
    )
    document.add_paragraph(
        "Before sending evidence, validate that the reviewed row count matches the original "
        "export, the sum of variance_amount equals the LedgerHub summary, every status is "
        "Clear or Pending, no required case ID is blank, and review_log.csv contains reviewer "
        "initials and a timestamp."
    )

    _add_labelled_paragraph(
        document,
        "WARNING — ",
        "Never edit the original CSV or overwrite a locked-period record. Work only in a "
        "copy and retain the untouched source.",
    )
    _add_labelled_paragraph(
        document,
        "NOTE — ",
        "Keep negative variance signs in the workbook; a negative amount indicates the invoice "
        "is below the purchase order.",
    )

    document.add_heading("Evidence and expected outputs", level=1)
    _add_bullet(
        document, "PO_Variance_YYYYMMDD_reviewed.xlsx with reason codes and Clear/Pending statuses."
    )
    _add_bullet(
        document,
        "Exception Queue case export with one case ID for every Pending row that was created "
        "successfully.",
    )
    _add_bullet(document, "review_log.csv containing reviewer initials and the review timestamp.")
    _add_bullet(
        document,
        "variance_summary_YYYYMMDD.png and attached invoice PDFs in the weekly Evidence folder.",
    )
    _add_bullet(
        document, "A Teams handoff containing the Finance\\AP\\Variance\\YYYY-WW\\Evidence link."
    )
    document.add_paragraph(
        "The procedure is complete only when the reviewed workbook, validation log, case evidence, "
        "and handoff link are present and every unresolved row is visibly Pending."
    )

    document.add_heading("Desk note copied from the older checklist", level=1)
    document.add_paragraph(
        "OLD DESK NOTE — Use an absolute total variance at or below $30 as Clear."
    )
    document.add_paragraph(
        "This $30 threshold conflicts with the $25 threshold in the LedgerHub decision panel and "
        "Step 6. The procedure does not identify which business owner can resolve the conflict; "
        "do not invent a value."
    )

    document.add_heading("Exceptions, recovery, and escalation", level=1)
    _add_labelled_paragraph(
        document,
        "Missing export: ",
        "Check the Procurement mailbox filter, refresh once, and look for a revised attachment. "
        "If the file is still missing, save a screenshot named missing_export_YYYYMMDD.png in "
        "the Evidence folder and notify the Procurement analyst in Teams.",
    )
    _add_labelled_paragraph(
        document,
        "LedgerHub or Exception Queue unavailable: ",
        "Wait 15 minutes and retry. Record the outage in review_log.csv, do not change statuses "
        "offline, and resume from the last saved workbook when access returns. If the service is "
        "not back by 14:00 ET, escalate to the finance owner referenced above; the owner and queue "
        "are not named in this checklist.",
    )
    _add_labelled_paragraph(
        document,
        "Validation failure: ",
        "Keep the affected row Pending, preserve the source row and invoice PDF, and create an "
        "Exception Queue case if the queue is available. Ask the Finance Reviewer to decide the "
        "classification before the evidence handoff.",
    )
    _add_labelled_paragraph(
        document,
        "Threshold conflict: ",
        "For a row between $25 and $30, treat the threshold as unresolved because Step 6 and the "
        "old desk note disagree. Leave the row Pending and ask the Finance Reviewer to route the "
        "question to the unnamed finance owner.",
    )
    document.add_paragraph(
        "Recovery is complete when the source is unchanged, the latest working copy is saved, "
        "the outage or conflict is recorded, and the Finance Reviewer has the evidence needed to "
        "continue or escalate."
    )

    # A short footer gives the source a realistic bit of page furniture without
    # changing the substantive facts extracted by the fixture tests.
    for footer_section in document.sections:
        footer = footer_section.footer.paragraphs[0]
        footer.alignment = 2
        footer.add_run("Weekly PO Invoice Variance Review | Internal synthetic fixture")

    document.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_fixture())
