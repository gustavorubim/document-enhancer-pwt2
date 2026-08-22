import hashlib
import json
import re
from pathlib import Path

from docx import Document
from PIL import Image
from typer.testing import CliRunner

from document_enhancer.cli import app

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "tests" / "fixtures" / "messy_desktop_procedure.docx"
TEMPLATE = ROOT / "templates" / "desktop_procedure.md"
STAGE1_FILES = {
    "draft.md",
    "draft.docx",
    "analysis.md",
    "analysis.docx",
    "mapping.json",
    "process_flow.mmd",
    "process_flow.png",
    "questions.json",
}
STAGE2_FILES = {
    "final.md",
    "final.docx",
    "resolution.json",
    "process_flow.mmd",
    "process_flow.png",
}


def _stage1(tmp_path: Path) -> Path:
    output = tmp_path / "stage1"
    result = CliRunner().invoke(
        app,
        [
            "stage1",
            "--source",
            str(SOURCE),
            "--template",
            str(TEMPLATE),
            "--output-dir",
            str(output),
            "--provider",
            "fake",
        ],
    )
    assert result.exit_code == 0, result.output
    return output


def _complete_answers(path: Path) -> None:
    payload = json.loads(path.read_text(encoding="utf-8"))
    for question in payload["questions"]:
        if "threshold" in question["text"].casefold():
            question["answer"] = (
                "Use an absolute total variance threshold at or below $25 for Clear."
            )
        else:
            question["answer"] = "Use the Finance Operations escalation queue in Teams."
    path.write_text(json.dumps(payload, indent=2) + "\n", encoding="utf-8")


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_stage1_adds_source_grounded_process_flow_and_editable_questions(
    tmp_path: Path,
) -> None:
    output = _stage1(tmp_path)

    assert {path.name for path in output.iterdir()} == STAGE1_FILES
    mermaid = (output / "process_flow.mmd").read_text(encoding="utf-8")
    draft = (output / "draft.md").read_text(encoding="utf-8")
    questions = json.loads((output / "questions.json").read_text(encoding="utf-8"))
    assert mermaid.startswith("flowchart TD\n")
    assert mermaid.count("NODE_") >= 12
    assert "Pending / exception" in mermaid
    assert "![Source-derived process flow](process_flow.png)" in draft
    assert "```mermaid" in draft and "flowchart TD" in draft
    assert len(questions["questions"]) == 2
    assert all(question["answer"] == "" for question in questions["questions"])
    with Image.open(output / "process_flow.png") as image:
        assert image.width >= 900 and image.height >= 800
    document = Document(output / "draft.docx")
    assert len(document.inline_shapes) == 1
    assert "Mermaid source" in [paragraph.text for paragraph in document.paragraphs]
    assert any("flowchart TD" in paragraph.text for paragraph in document.paragraphs)


def test_stage2_uses_answers_without_modifying_stage1(tmp_path: Path) -> None:
    stage1 = _stage1(tmp_path)
    answers = stage1 / "questions.json"
    _complete_answers(answers)
    protected = {
        path.name: _sha256(path) for path in stage1.iterdir() if path.name != "questions.json"
    }
    final_dir = tmp_path / "stage2"

    result = CliRunner().invoke(
        app,
        [
            "stage2",
            "--source",
            str(SOURCE),
            "--template",
            str(TEMPLATE),
            "--answers",
            str(answers),
            "--output-dir",
            str(final_dir),
            "--provider",
            "fake",
        ],
    )

    assert result.exit_code == 0, result.output
    assert {path.name for path in final_dir.iterdir()} == STAGE2_FILES
    final = (final_dir / "final.md").read_text(encoding="utf-8")
    assert "[MISSING:" not in final
    assert "[CONFLICT:" not in final
    assert "$30" not in final
    assert len(re.findall(r"(?m)^\d+\. ", final.split("## Process flow", maxsplit=1)[0])) == 9
    assert "Use an absolute total variance threshold at or below $25 for Clear." in final
    assert "Use the Finance Operations escalation queue in Teams." in final
    body = final.split("## Process flow", maxsplit=1)[0]
    assert body.count("Use an absolute total variance threshold at or below $25 for Clear.") == 1
    assert body.count("Use the Finance Operations escalation queue in Teams.") == 2
    assert "![Source-derived process flow](process_flow.png)" in final
    resolution = json.loads((final_dir / "resolution.json").read_text(encoding="utf-8"))
    assert len(resolution["resolutions"]) == 2
    assert all(
        item["answer_source_block_id"].startswith("SRC-") for item in resolution["resolutions"]
    )
    assert len(Document(final_dir / "final.docx").inline_shapes) == 1
    assert protected == {
        path.name: _sha256(path) for path in stage1.iterdir() if path.name != "questions.json"
    }


def test_stage2_rejects_incomplete_answers(tmp_path: Path) -> None:
    stage1 = _stage1(tmp_path)
    result = CliRunner().invoke(
        app,
        [
            "stage2",
            "--source",
            str(SOURCE),
            "--template",
            str(TEMPLATE),
            "--answers",
            str(stage1 / "questions.json"),
            "--output-dir",
            str(tmp_path / "stage2"),
            "--provider",
            "fake",
        ],
    )

    assert result.exit_code == 2
    assert "Stage 2 answer is empty" in result.output
