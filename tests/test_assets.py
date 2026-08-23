import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw

from document_enhancer.ingest import ingest_source
from document_enhancer.models import TemplateRequirement, TemplateSection
from document_enhancer.pipeline import (
    DeterministicProvider,
    _insert_after_anchor_in_section,
    run_enhancement,
)

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "desktop_procedure.md"


def _screenshot(path: Path, *, color: str = "#DCEBFF") -> bytes:
    image = Image.new("RGB", (480, 260), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((20, 20, 460, 240), radius=16, fill=color, outline="#2255AA", width=4)
    draw.rectangle((45, 55, 435, 95), fill="#2255AA")
    draw.text((60, 65), "Atlas Request Review", fill="white")
    draw.rectangle((55, 125, 260, 165), fill="white", outline="#667788", width=2)
    draw.rectangle((300, 188, 415, 222), fill="#2B7A3D")
    draw.text((330, 196), "Submit", fill="white")
    image.save(path)
    return path.read_bytes()


def _source_with_screenshot(tmp_path: Path) -> tuple[Path, bytes]:
    image_path = tmp_path / "atlas.png"
    payload = _screenshot(image_path)
    source_path = tmp_path / "procedure.docx"
    document = Document()
    document.add_heading("Atlas Request Review", level=0)
    document.add_heading("Purpose", level=1)
    document.add_paragraph("Review incoming requests before approval.")
    document.add_heading("Procedure steps", level=1)
    document.add_paragraph("Open Atlas and select the oldest request.")
    document.add_picture(str(image_path), width=Inches(4.2))
    document.add_paragraph("Select Submit and record the confirmation identifier.")
    document.add_heading("Evidence", level=1)
    document.add_paragraph("Save the confirmation identifier in review_log.csv.")
    document.save(source_path)
    return source_path, payload


def test_docx_ingestion_extracts_original_screenshot_bytes_and_anchor(tmp_path: Path) -> None:
    source_path, payload = _source_with_screenshot(tmp_path)

    source = ingest_source(source_path)

    assert len(source.assets) == 1
    asset = source.assets[0]
    assert asset.id == "FIG-001"
    assert asset.media_type == "image/png"
    assert asset.payload == payload
    assert asset.sha256 == hashlib.sha256(payload).hexdigest()
    assert asset.anchor_text == "Open Atlas and select the oldest request."
    assert asset.source_block_id in {block.id for block in source.blocks}


def test_pipeline_preserves_screenshot_in_markdown_mapping_and_docx(tmp_path: Path) -> None:
    source_path, payload = _source_with_screenshot(tmp_path)
    output_dir = tmp_path / "output"

    artifacts = run_enhancement(
        source_path=source_path,
        template_path=TEMPLATE,
        output_dir=output_dir,
        provider=DeterministicProvider(),
    )

    assert len(artifacts.source_asset_paths) == 1
    assert artifacts.source_asset_paths[0].read_bytes() == payload
    markdown = artifacts.draft_markdown.read_text(encoding="utf-8")
    assert "![Original source screenshot FIG-001](assets/FIG-001.png)" in markdown
    assert markdown.index("Open Atlas and select the oldest request.") < markdown.index(
        "assets/FIG-001.png"
    )
    mapping = json.loads(artifacts.mapping_json.read_text(encoding="utf-8"))
    assert mapping["source_assets"][0]["id"] == "FIG-001"
    assert mapping["source_assets"][0]["sha256"] == hashlib.sha256(payload).hexdigest()
    with zipfile.ZipFile(artifacts.draft_docx) as archive:
        embedded = [
            archive.read(name) for name in archive.namelist() if name.startswith("word/media/")
        ]
    assert payload in embedded


def test_docx_ingestion_preserves_screenshot_inside_table_cell(tmp_path: Path) -> None:
    image_path = tmp_path / "table-screenshot.png"
    payload = _screenshot(image_path, color="#F3E8FF")
    source_path = tmp_path / "table-procedure.docx"
    document = Document()
    document.add_heading("Table Procedure", level=0)
    document.add_heading("Procedure steps", level=1)
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.add_paragraph("Review the embedded request screenshot.")
    cell.add_paragraph().add_run().add_picture(str(image_path), width=Inches(3.5))
    document.save(source_path)

    source = ingest_source(source_path)

    assert len(source.assets) == 1
    assert source.assets[0].payload == payload
    assert source.assets[0].anchor_text == "Review the embedded request screenshot."


def test_duplicate_anchor_is_placed_in_mapped_target_section() -> None:
    markdown = """# Procedure

## First section

Repeat this instruction.

## Second section

Repeat this instruction.
"""
    target = TemplateSection(
        id="SEC-002",
        heading="Second section",
        level=2,
        requirements=[TemplateRequirement(id="REQ-002-01", text="Retain the screenshot.")],
    )

    placed = _insert_after_anchor_in_section(
        markdown,
        target,
        "Repeat this instruction.",
        "![FIG-001](assets/FIG-001.png)",
    )

    assert placed is not None
    first, second = placed.split("## Second section", maxsplit=1)
    assert "FIG-001" not in first
    assert "Repeat this instruction.\n\n![FIG-001]" in second
