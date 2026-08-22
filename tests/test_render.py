import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image

from document_enhancer.render import (
    RenderError,
    render_markdown_file,
    render_markdown_pair,
    render_markdown_to_docx,
)


def _paragraph_num_id(paragraph) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return None
    num_id = p_pr.numPr.find(qn("w:numId"))
    return None if num_id is None else num_id.get(qn("w:val"))


def test_render_preserves_sections_inline_formatting_lists_and_table(tmp_path: Path) -> None:
    markdown = """# Draft procedure

Intro with **bold**, *emphasis*, and `access.csv`.

## Actions

1. Open the review queue.
2. Save the result as `access.csv`.

- Confirm the file exists.
- Escalate if the export fails.

| Requirement | Status |
| --- | --- |
| Purpose | Supported |
"""
    output = tmp_path / "draft.docx"

    assert render_markdown_to_docx(markdown, output) == output

    document = Document(output)
    assert [paragraph.text for paragraph in document.paragraphs[:3]] == [
        "Draft procedure",
        "Intro with bold, emphasis, and access.csv.",
        "Actions",
    ]
    assert [paragraph.style.name for paragraph in document.paragraphs[:3]] == [
        "Heading 1",
        "Normal",
        "Heading 2",
    ]
    list_paragraphs = document.paragraphs[3:]
    assert [paragraph.text for paragraph in list_paragraphs] == [
        "Open the review queue.",
        "Save the result as access.csv.",
        "Confirm the file exists.",
        "Escalate if the export fails.",
    ]
    assert [paragraph.style.name for paragraph in list_paragraphs] == [
        "List Number",
        "List Number",
        "List Bullet",
        "List Bullet",
    ]
    assert all(_paragraph_num_id(paragraph) is not None for paragraph in list_paragraphs)

    table = document.tables[0]
    assert [[cell.text for cell in row.cells] for row in table.rows] == [
        ["Requirement", "Status"],
        ["Purpose", "Supported"],
    ]
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    table_indent = table._tbl.tblPr.find(qn("w:tblInd"))
    assert table_width is not None and table_width.get(qn("w:w")) == "9360"
    assert table_indent is not None and table_indent.get(qn("w:w")) == "120"


def test_render_file_pair_writes_matching_draft_and_analysis_names(tmp_path: Path) -> None:
    draft = tmp_path / "draft.md"
    analysis = tmp_path / "analysis.md"
    draft.write_text("# Draft\n\n## Purpose\n\nSource-grounded content.", encoding="utf-8")
    analysis.write_text("# Analysis\n\n## Coverage\n\nSupported.", encoding="utf-8")

    outputs = render_markdown_pair(draft, analysis, tmp_path / "rendered")

    assert outputs.draft_docx.name == "draft.docx"
    assert outputs.analysis_docx.name == "analysis.docx"
    assert [paragraph.text for paragraph in Document(outputs.draft_docx).paragraphs] == [
        "Draft",
        "Purpose",
        "Source-grounded content.",
    ]
    assert [paragraph.text for paragraph in Document(outputs.analysis_docx).paragraphs] == [
        "Analysis",
        "Coverage",
        "Supported.",
    ]


def test_render_file_embeds_standalone_local_image_with_alt_text(tmp_path: Path) -> None:
    image_path = tmp_path / "process_flow.png"
    Image.new("RGB", (1000, 250), color=(34, 116, 181)).save(image_path)
    markdown_path = tmp_path / "procedure.md"
    markdown_path.write_text(
        "# Procedure\n\nBefore the image.\n\n"
        "![Process flow](process_flow.png)\n\nAfter the image.\n",
        encoding="utf-8",
    )
    output = tmp_path / "procedure.docx"

    assert render_markdown_file(markdown_path, output) == output

    document = Document(output)
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Procedure",
        "Before the image.",
        "",
        "After the image.",
    ]
    assert len(document.inline_shapes) == 1
    shape = document.inline_shapes[0]
    assert shape.width <= Inches(6.5)
    assert document.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert shape._inline.docPr.get("descr") == "Process flow"

    with zipfile.ZipFile(output) as archive:
        media_parts = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(media_parts) == 1


@pytest.mark.parametrize(
    ("markdown", "expected_message"),
    [
        ("![Missing](missing.png)", "does not exist"),
        ("![Remote](https://example.com/process.png)", "remote image URLs"),
    ],
)
def test_render_file_rejects_unavailable_images(
    tmp_path: Path, markdown: str, expected_message: str
) -> None:
    markdown_path = tmp_path / "procedure.md"
    markdown_path.write_text(markdown, encoding="utf-8")

    with pytest.raises(RenderError, match=expected_message):
        render_markdown_file(markdown_path, tmp_path / "procedure.docx")


def test_render_file_rejects_unsupported_image_format(tmp_path: Path) -> None:
    image_path = tmp_path / "not-an-image.txt"
    image_path.write_text("not an image", encoding="utf-8")
    markdown_path = tmp_path / "procedure.md"
    markdown_path.write_text("![Unsupported](not-an-image.txt)", encoding="utf-8")

    with pytest.raises(RenderError, match="unsupported or unreadable"):
        render_markdown_file(markdown_path, tmp_path / "procedure.docx")


@pytest.mark.parametrize(
    ("markdown", "output_name", "message"),
    [
        ("", "draft.docx", "empty"),
        ("# Draft", "draft.txt", "suffix"),
    ],
)
def test_render_rejects_invalid_artifacts(
    tmp_path: Path, markdown: str, output_name: str, message: str
) -> None:
    with pytest.raises(RenderError, match=message):
        render_markdown_to_docx(markdown, tmp_path / output_name)


def test_render_file_rejects_missing_markdown(tmp_path: Path) -> None:
    with pytest.raises(RenderError, match="does not exist"):
        render_markdown_file(tmp_path / "missing.md", tmp_path / "draft.docx")
