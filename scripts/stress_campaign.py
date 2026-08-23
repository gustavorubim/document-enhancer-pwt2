"""Generate and execute the deterministic 20-document stress campaign."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from PIL import Image, ImageDraw

from document_enhancer.batch import SUPPORTED_SUFFIXES, run_batch
from document_enhancer.pipeline import DeterministicProvider


@dataclass(frozen=True)
class CampaignSpec:
    name: str
    pages: int
    profile: str
    screenshot: bool
    table: bool
    conflict: bool
    missing_owner: bool


PAGE_COUNTS = (5, 6, 7, 8, 9, 10, 11, 12, 14, 15, 16, 18, 20, 22, 24, 26, 27, 28, 29, 30)
SECTION_NAMES = (
    "Purpose and scope",
    "Prerequisites and access",
    "Roles and timing",
    "Procedure steps",
    "Decision points and validation",
    "Outputs and evidence",
    "Exceptions and recovery",
)


def campaign_specs() -> list[CampaignSpec]:
    profiles = ("structured", "unstructured", "mixed")
    return [
        CampaignSpec(
            name=f"campaign-{index:02d}-{profiles[(index - 1) % len(profiles)]}",
            pages=pages,
            profile=profiles[(index - 1) % len(profiles)],
            screenshot=index % 3 == 0,
            table=index % 4 == 0,
            conflict=index % 5 == 0,
            missing_owner=index % 6 == 0,
        )
        for index, pages in enumerate(PAGE_COUNTS, start=1)
    ]


def generate_campaign(input_dir: Path) -> list[CampaignSpec]:
    input_dir.mkdir(parents=True, exist_ok=True)
    image_dir = input_dir / "_screenshots"
    image_dir.mkdir(exist_ok=True)
    specs = campaign_specs()
    for index, spec in enumerate(specs, start=1):
        screenshot_path = image_dir / f"{spec.name}.png"
        if spec.screenshot:
            _build_screenshot(screenshot_path, index)
        _build_document(input_dir / f"{spec.name}.docx", spec, screenshot_path)
    (input_dir / "campaign_spec.json").write_text(
        json.dumps([asdict(spec) for spec in specs], indent=2) + "\n",
        encoding="utf-8",
    )
    return specs


def declared_page_count(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return 1 + len(re.findall(r'<w:br[^>]*w:type="page"', document_xml))


def run_campaign(*, work_dir: Path, template_path: Path) -> dict[str, object]:
    input_dir = work_dir / "input"
    output_dir = work_dir / "output"
    expected_sources = {f"{spec.name}.docx" for spec in campaign_specs()}
    unexpected_sources = sorted(
        path.name
        for path in input_dir.glob("*")
        if path.is_file()
        and path.suffix.casefold() in SUPPORTED_SUFFIXES
        and path.name not in expected_sources
    )
    if unexpected_sources:
        names = ", ".join(unexpected_sources)
        raise RuntimeError(f"campaign input contains unexpected source files: {names}")
    specs = generate_campaign(input_dir)
    manifest = run_batch(
        input_dir=input_dir,
        template_path=template_path,
        output_dir=output_dir,
        provider=DeterministicProvider(),
    )
    page_counts = {
        spec.name: declared_page_count(input_dir / f"{spec.name}.docx") for spec in specs
    }
    expected_screenshot_documents = sum(spec.screenshot for spec in specs)
    verified_screenshot_documents = sum(
        _screenshot_preserved(input_dir=input_dir, output_dir=output_dir, spec=spec)
        for spec in specs
        if spec.screenshot
    )
    expected_recovery_minimum = sum(spec.profile == "unstructured" for spec in specs)
    report = {
        "schema_version": "stress-campaign.v1",
        "document_count": len(specs),
        "minimum_pages": min(page_counts.values()),
        "maximum_pages": max(page_counts.values()),
        "page_count_method": "explicit_ooxml_page_breaks",
        "declared_pages": page_counts,
        "expected_screenshot_documents": expected_screenshot_documents,
        "verified_screenshot_documents": verified_screenshot_documents,
        "expected_recovery_minimum": expected_recovery_minimum,
        "manifest": manifest.model_dump(mode="json"),
    }
    (work_dir / "stress_report.json").write_text(
        json.dumps(report, indent=2) + "\n", encoding="utf-8"
    )
    if manifest.failed_count:
        raise RuntimeError(f"stress campaign had {manifest.failed_count} failed documents")
    if (
        len(specs) != 20
        or len(manifest.documents) != 20
        or min(page_counts.values()) != 5
        or max(page_counts.values()) != 30
    ):
        raise RuntimeError("stress campaign page-range contract was not satisfied")
    if verified_screenshot_documents != expected_screenshot_documents:
        raise RuntimeError("stress campaign did not preserve every source screenshot")
    if manifest.recovered_count < expected_recovery_minimum:
        raise RuntimeError("stress campaign did not recover every unstructured document")
    return report


def _screenshot_preserved(*, input_dir: Path, output_dir: Path, spec: CampaignSpec) -> bool:
    source_payload = (input_dir / "_screenshots" / f"{spec.name}.png").read_bytes()
    document_output = output_dir / spec.name
    asset_path = document_output / "assets" / "FIG-001.png"
    if not asset_path.is_file() or asset_path.read_bytes() != source_payload:
        return False
    with zipfile.ZipFile(document_output / "draft.docx") as archive:
        embedded_media = [
            archive.read(name) for name in archive.namelist() if name.startswith("word/media/")
        ]
    return source_payload in embedded_media


def _build_screenshot(path: Path, index: int) -> None:
    image = Image.new("RGB", (640, 360), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle(
        (24, 24, 616, 336), radius=18, fill="#EAF2FF", outline="#2457A6", width=4
    )
    draw.rectangle((50, 54, 590, 108), fill="#2457A6")
    draw.text((72, 72), f"Campaign request review {index:02d}", fill="white")
    draw.rectangle((70, 145, 360, 195), fill="white", outline="#667788", width=2)
    draw.rectangle((420, 260, 565, 310), fill="#2F7D45")
    draw.text((460, 275), "Submit", fill="white")
    image.save(path)


def _build_document(path: Path, spec: CampaignSpec, screenshot_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    document.add_heading(f"Stress Procedure {spec.name}", level=0)
    for page_index in range(1, spec.pages + 1):
        section_name = SECTION_NAMES[(page_index - 1) % len(SECTION_NAMES)]
        if spec.profile == "structured" or (spec.profile == "mixed" and page_index % 3 == 1):
            document.add_heading(section_name, level=1)
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{section_name}: ").bold = True
            paragraph.add_run("operational notes continue below without a Word heading style.")
        _add_page_content(document, spec, page_index)
        if spec.table and page_index == 2:
            table = document.add_table(rows=3, cols=3)
            for column, value in enumerate(("Check", "Owner", "Evidence")):
                table.cell(0, column).text = value
            table.cell(1, 0).text = "Request identifier present"
            table.cell(1, 1).text = "Operations Analyst"
            table.cell(1, 2).text = "review_log.csv"
            table.cell(2, 0).text = "Confirmation identifier saved"
            table.cell(2, 1).text = "Finance Reviewer"
            table.cell(2, 2).text = "Atlas confirmation"
        if spec.screenshot and page_index == 2:
            document.add_paragraph("Confirm the Atlas request screen before selecting Submit.")
            document.add_picture(str(screenshot_path), width=Inches(4.7))
        if page_index < spec.pages:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.save(path)


def _add_page_content(document: Document, spec: CampaignSpec, page_index: int) -> None:
    document.add_paragraph(
        f"Page {page_index} supports the weekly request review performed by the Operations Analyst "
        "after the request export arrives."
    )
    document.add_paragraph(
        "Atlas access, approved_accounts.csv, and a writable evidence folder are required "
        "before work begins."
    )
    document.add_paragraph(
        f"Open Atlas, select request {spec.name.upper()}-{page_index:03d}, compare the account "
        "identifier, and keep unsupported requests on Hold."
    )
    document.add_paragraph(
        "Validate the status, save the confirmation identifier in review_log.csv, and "
        "retain the evidence link."
    )
    if spec.conflict and page_index in {2, 3}:
        threshold = "$25" if page_index == 2 else "$30"
        document.add_paragraph(
            f"Use a total variance threshold at or below {threshold} for Clear; the two "
            "source pages conflict."
        )
    if spec.missing_owner and page_index == 3:
        document.add_paragraph(
            "Escalate an outage after 15 minutes, but the source does not identify the "
            "escalation owner or queue."
        )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--work-dir", type=Path, default=Path("runs/stress-campaign"))
    parser.add_argument("--template", type=Path, default=Path("templates/desktop_procedure.md"))
    args = parser.parse_args()
    report = run_campaign(work_dir=args.work_dir, template_path=args.template)
    print(args.work_dir / "stress_report.json")
    print(
        f"documents={report['document_count']} pages={report['minimum_pages']}-"
        f"{report['maximum_pages']}"
    )


if __name__ == "__main__":
    main()
