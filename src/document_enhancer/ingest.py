"""Normalize complete Markdown or DOCX sources into ordered, traceable blocks."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from markdown_it import MarkdownIt
from PIL import Image as PillowImage

from document_enhancer.models import SourceAsset, SourceBlock, SourceDocument, SourceFormat


class SourceIngestionError(ValueError):
    """Raised when a source cannot be normalized safely."""


@dataclass(frozen=True)
class _RawBlock:
    heading: str
    content: str


def ingest_source(path: Path) -> SourceDocument:
    """Read an entire supported source and return stable, ordered source blocks."""

    source_path = path.expanduser().resolve()
    if not source_path.is_file():
        raise SourceIngestionError(f"source file does not exist: {source_path}")

    suffix = source_path.suffix.casefold()
    if suffix in {".md", ".markdown"}:
        return _ingest_markdown(source_path)
    if suffix == ".docx":
        return _ingest_docx(source_path)
    raise SourceIngestionError(
        f"unsupported source format {source_path.suffix!r}; expected .docx, .md, or .markdown"
    )


def _ingest_markdown(path: Path) -> SourceDocument:
    raw = path.read_text(encoding="utf-8").strip()
    if not raw:
        raise SourceIngestionError(f"source document is empty: {path}")

    lines = raw.splitlines()
    tokens = MarkdownIt("commonmark").parse(raw)
    headings: list[tuple[int, str, int]] = []
    for index, token in enumerate(tokens):
        if token.type != "heading_open" or token.map is None:
            continue
        inline = tokens[index + 1]
        headings.append((token.map[0], inline.content.strip(), int(token.tag[1])))

    title = next((heading for _, heading, level in headings if level == 1), path.stem)
    raw_blocks = _markdown_sections(lines, headings)
    return _source_document(path, SourceFormat.MARKDOWN, title, raw_blocks, raw)


def _markdown_sections(lines: list[str], headings: list[tuple[int, str, int]]) -> list[_RawBlock]:
    if not headings:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", "\n".join(lines))]
        return [
            _RawBlock(heading=f"Source content {index}", content=content)
            for index, content in enumerate(paragraphs, start=1)
            if _is_meaningful(content)
        ]

    blocks: list[_RawBlock] = []
    preamble = "\n".join(lines[: headings[0][0]]).strip()
    if _is_meaningful(preamble):
        blocks.append(_RawBlock(heading="Preamble", content=preamble))

    for index, (line_number, heading, _level) in enumerate(headings):
        end = headings[index + 1][0] if index + 1 < len(headings) else len(lines)
        content = "\n".join(lines[line_number + 1 : end]).strip()
        if _is_meaningful(content):
            blocks.append(_RawBlock(heading=heading, content=content))
    return blocks


def _ingest_docx(path: Path) -> SourceDocument:
    try:
        document = Document(path)
    except Exception as exc:  # python-docx exposes several package/XML errors
        raise SourceIngestionError(f"could not read DOCX source {path}: {exc}") from exc

    title, raw_blocks = _docx_sections(document, path.stem)
    if len(raw_blocks) == 1 and raw_blocks[0].heading in {"Preamble", title}:
        parts = [part.strip() for part in re.split(r"\n\s*\n", raw_blocks[0].content)]
        meaningful_parts = [part for part in parts if _is_meaningful(part)]
        if len(meaningful_parts) > 1:
            raw_blocks = [
                _RawBlock(heading=f"Source content {index}", content=content)
                for index, content in enumerate(meaningful_parts, start=1)
            ]
    if not raw_blocks:
        raise SourceIngestionError(f"source document has no meaningful text: {path}")
    full_text = "\n\n".join(f"{block.heading}\n{block.content}" for block in raw_blocks)
    source = _source_document(path, SourceFormat.DOCX, title, raw_blocks, full_text)
    assets = _extract_docx_assets(document, source.blocks)
    return SourceDocument(
        title=source.title,
        source_path=source.source_path,
        source_format=source.source_format,
        blocks=source.blocks,
        assets=assets,
        full_text=source.full_text,
    )


def _extract_docx_assets(document: DocumentObject, blocks: list[SourceBlock]) -> list[SourceAsset]:
    assets: list[SourceAsset] = []
    current_heading = "Preamble"
    last_text = ""
    for item in _iter_body_paragraphs(document):
        text = item.text.strip()
        style_name = (item.style.name if item.style else "").strip()
        if style_name.casefold() == "title" or _heading_level(style_name) is not None:
            current_heading = text or current_heading
        blips = item._p.xpath(".//a:blip")
        descriptions = [
            str(node.get("descr") or node.get("title") or "").strip()
            for node in item._p.xpath(".//wp:docPr")
        ]
        for occurrence, blip in enumerate(blips):
            relation_id = blip.get(qn("r:embed"))
            if not relation_id:
                continue
            part = document.part.related_parts.get(relation_id)
            media_type = getattr(part, "content_type", "")
            if media_type not in {"image/png", "image/jpeg"}:
                continue
            payload = bytes(getattr(part, "blob", b""))
            if not payload:
                continue
            try:
                with PillowImage.open(BytesIO(payload)) as image:
                    width, height = image.size
            except OSError as exc:
                raise SourceIngestionError(
                    f"embedded DOCX image {relation_id} is unreadable"
                ) from exc
            anchor_text = text or last_text
            source_block_id = _asset_source_block_id(
                blocks, heading=current_heading, anchor_text=anchor_text
            )
            digest = hashlib.sha256(payload).hexdigest()
            partname = str(getattr(part, "partname", relation_id))
            figure_id = f"FIG-{len(assets) + 1:03d}"
            alt_text = descriptions[occurrence] if occurrence < len(descriptions) else ""
            assets.append(
                SourceAsset(
                    id=figure_id,
                    source_block_id=source_block_id,
                    order=len(assets) + 1,
                    original_name=Path(partname).name or f"{figure_id}.png",
                    media_type=media_type,
                    sha256=digest,
                    size_bytes=len(payload),
                    width_pixels=width,
                    height_pixels=height,
                    anchor_text=anchor_text[:500],
                    alt_text=alt_text,
                    payload=payload,
                )
            )
        if text and not blips:
            last_text = text
    return assets


def _iter_body_paragraphs(document: DocumentObject):
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            yield item
        elif isinstance(item, Table):
            yield from _iter_table_paragraphs(item)


def _iter_table_paragraphs(table: Table):
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            for item in cell.iter_inner_content():
                if isinstance(item, Paragraph):
                    yield item
                elif isinstance(item, Table):
                    yield from _iter_table_paragraphs(item)


def _asset_source_block_id(blocks: list[SourceBlock], *, heading: str, anchor_text: str) -> str:
    normalized_anchor = " ".join(anchor_text.split())
    if normalized_anchor:
        for block in blocks:
            if normalized_anchor in " ".join(block.content.split()):
                return block.id
    normalized_heading = heading.strip().casefold()
    for block in blocks:
        if block.heading.strip().casefold() == normalized_heading:
            return block.id
    return blocks[0].id


def _docx_sections(document: DocumentObject, fallback_title: str) -> tuple[str, list[_RawBlock]]:
    title = fallback_title
    blocks: list[_RawBlock] = []
    current_heading = "Preamble"
    current_content: list[str] = []
    list_number = 0

    def flush() -> None:
        content = "\n\n".join(current_content).strip()
        if _is_meaningful(content):
            blocks.append(_RawBlock(heading=current_heading, content=content))
        current_content.clear()

    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            style_name = (item.style.name if item.style else "").strip()
            heading_level = _heading_level(style_name)
            if style_name.casefold() == "title":
                title = text
                flush()
                current_heading = text
                list_number = 0
                continue
            if heading_level is not None:
                if heading_level == 1 and title == fallback_title:
                    title = text
                flush()
                current_heading = text
                list_number = 0
                continue

            formatted, list_number = _format_docx_paragraph(item, text, list_number)
            current_content.append(formatted)
        elif isinstance(item, Table):
            table_markdown = _table_as_markdown(item)
            if table_markdown:
                current_content.append(table_markdown)
            list_number = 0

    flush()
    return title, blocks


def _heading_level(style_name: str) -> int | None:
    match = re.fullmatch(r"Heading\s+([1-6])", style_name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def _format_docx_paragraph(paragraph: Paragraph, text: str, counter: int) -> tuple[str, int]:
    style_name = (paragraph.style.name if paragraph.style else "").casefold()
    if "list number" in style_name:
        counter += 1
        return f"{counter}. {text}", counter
    if "list bullet" in style_name:
        return f"- {text}", 0
    return text, 0


def _table_as_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()).replace("|", "\\|") for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = f"| {' | '.join(padded[0])} |"
    separator = f"| {' | '.join(['---'] * width)} |"
    body = [f"| {' | '.join(row)} |" for row in padded[1:]]
    return "\n".join([header, separator, *body])


def _source_document(
    path: Path,
    source_format: SourceFormat,
    title: str,
    raw_blocks: list[_RawBlock],
    full_text: str,
) -> SourceDocument:
    if not raw_blocks:
        raise SourceIngestionError(f"source document has no meaningful sections: {path}")
    blocks = [
        SourceBlock(
            id=f"SRC-{index:03d}",
            heading=block.heading,
            content=block.content,
            order=index,
        )
        for index, block in enumerate(raw_blocks, start=1)
    ]
    return SourceDocument(
        title=title.strip() or path.stem,
        source_path=path,
        source_format=source_format,
        blocks=blocks,
        full_text=full_text,
    )


def _is_meaningful(text: str) -> bool:
    return bool(re.search(r"[\w\d]", text, flags=re.UNICODE))
