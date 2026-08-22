"""Render source-grounded Markdown artifacts as deterministic, styled DOCX files.

The renderer intentionally operates on Markdown rather than on provider-specific
models.  This keeps the output contract small: callers write the ordered
``draft.md`` or ``analysis.md`` artifact and pass it to
:func:`render_markdown_file`.  The same Markdown is parsed for both the human
readable and DOCX outputs, so headings, lists, tables, and inline formatting
cannot silently drift between them.

Only a compact subset of Markdown is needed by the MVP, but the implementation
uses ``markdown-it-py`` tokens instead of line-oriented heuristics.  In
particular, lists are written with real Word numbering definitions and tables
carry explicit DXA geometry so that wrapped content remains readable in Word
and LibreOffice.
"""

from __future__ import annotations

from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token

__all__ = [
    "RenderError",
    "RenderedDocuments",
    "render_document",
    "render_markdown",
    "render_markdown_file",
    "render_markdown_pair",
    "render_markdown_text",
    "render_markdown_to_docx",
]


# compact_reference_guide preset tokens.  Keep these as explicit constants so
# the generated document does not depend on Word's Normal style or page defaults.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120

BODY_FONT = "Calibri"
BODY_SIZE_PT = 11.0
BODY_AFTER_PT = 6.0
BODY_LINE_SPACING = 1.25
HEADING_COLOR = "2E74B5"
HEADING_DARK_COLOR = "1F4D78"
MUTED_COLOR = "667085"
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER_COLOR = "B8C2CC"
CODE_FILL = "F2F4F7"

LIST_MARKER_AT_IN = 0.187
LIST_TEXT_AT_IN = 0.375
LIST_HANGING_IN = 0.188
LIST_LEVEL_STEP_IN = 0.25
LIST_MAX_LEVEL = 8


class RenderError(ValueError):
    """Raised when Markdown cannot be rendered to a requested DOCX artifact."""


@dataclass(frozen=True)
class RenderedDocuments:
    """Paths produced by :func:`render_markdown_pair`."""

    draft_docx: Path
    analysis_docx: Path


@dataclass(frozen=True)
class _InlineCell:
    """A table cell's inline tokens and whether it is a header cell."""

    children: tuple[Token, ...]
    header: bool


def _parser() -> MarkdownIt:
    """Return the Markdown parser used by every renderer entry point."""

    # ``table`` is not part of the CommonMark preset in markdown-it-py.  Enable
    # it explicitly because analysis and draft artifacts commonly contain
    # requirement/status matrices.
    return MarkdownIt("commonmark").enable("table")


def render_markdown_file(
    markdown_path: Path | str,
    docx_path: Path | str,
    *,
    title: str | None = None,
) -> Path:
    """Render a UTF-8 Markdown file to ``docx_path`` and return that path.

    ``markdown_path`` and ``docx_path`` are resolved only after validation so a
    bad input produces a useful :class:`RenderError` instead of a low-level
    ``FileNotFoundError`` or a partially written DOCX.
    """

    source_path = Path(markdown_path).expanduser()
    if not source_path.is_file():
        raise RenderError(f"Markdown artifact does not exist: {source_path}")
    try:
        markdown = source_path.read_text(encoding="utf-8")
    except OSError as exc:
        raise RenderError(f"could not read Markdown artifact {source_path}: {exc}") from exc
    return render_markdown_to_docx(
        markdown,
        docx_path,
        title=title or _title_from_markdown(markdown),
        base_dir=source_path.parent,
    )


def render_markdown_text(
    markdown: str,
    docx_path: Path | str,
    *,
    title: str | None = None,
) -> Path:
    """Render Markdown text to a DOCX file and return the output path."""

    return render_markdown_to_docx(markdown, docx_path, title=title)


def render_markdown_to_docx(
    markdown: str | Path,
    docx_path: Path | str,
    *,
    title: str | None = None,
    base_dir: Path | str | None = None,
) -> Path:
    """Render Markdown text (or a ``Path`` to a Markdown file) to DOCX.

    The ``Path`` form is accepted as a convenience for pipeline callers.  Use
    :func:`render_markdown_file` when a string path should be unambiguous.
    """

    image_base_dir = Path(base_dir).expanduser() if base_dir is not None else None
    if isinstance(markdown, Path):
        source_path = markdown.expanduser()
        if not source_path.is_file():
            raise RenderError(f"Markdown artifact does not exist: {source_path}")
        try:
            markdown_text = source_path.read_text(encoding="utf-8")
        except OSError as exc:
            raise RenderError(f"could not read Markdown artifact {source_path}: {exc}") from exc
        if image_base_dir is None:
            image_base_dir = source_path.parent
    elif isinstance(markdown, str):
        # Accept a string path as a convenience while keeping ordinary Markdown
        # strings (which often contain a newline) unambiguous.
        candidate = Path(markdown).expanduser()
        try:
            is_file = "\n" not in markdown and candidate.is_file()
        except OSError:
            is_file = False
        if is_file:
            try:
                markdown_text = candidate.read_text(encoding="utf-8")
            except OSError as exc:
                raise RenderError(f"could not read Markdown artifact {candidate}: {exc}") from exc
            if image_base_dir is None:
                image_base_dir = candidate.parent
        else:
            markdown_text = markdown
    else:  # pragma: no cover - type checkers catch this; useful for dynamic callers.
        raise RenderError("Markdown input must be text or a pathlib.Path")

    if not markdown_text.strip():
        raise RenderError("Markdown artifact is empty")

    output_path = Path(docx_path).expanduser()
    if output_path.suffix.casefold() != ".docx":
        raise RenderError(f"DOCX output must have a .docx suffix: {output_path}")
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise RenderError(
            f"could not create DOCX output directory {output_path.parent}: {exc}"
        ) from exc

    document = _new_document(title or _title_from_markdown(markdown_text) or output_path.stem)
    tokens = _parser().parse(markdown_text)
    _render_blocks(document, tokens, base_dir=image_base_dir)
    _set_header_footer(document, title or _title_from_markdown(markdown_text) or output_path.stem)
    try:
        document.save(output_path)
    except OSError as exc:
        raise RenderError(f"could not write DOCX artifact {output_path}: {exc}") from exc
    return output_path


def render_markdown(
    markdown: str | Path,
    docx_path: Path | str,
    *,
    title: str | None = None,
) -> Path:
    """Compatibility alias for :func:`render_markdown_to_docx`."""

    return render_markdown_to_docx(markdown, docx_path, title=title)


def render_document(
    markdown_path: Path | str,
    docx_path: Path | str,
    *,
    title: str | None = None,
) -> Path:
    """Compatibility alias for rendering a Markdown document file."""

    return render_markdown_file(markdown_path, docx_path, title=title)


def render_markdown_pair(
    draft_markdown: Path | str,
    analysis_markdown: Path | str,
    output_dir: Path | str,
) -> RenderedDocuments:
    """Render the ordered draft and analysis Markdown artifacts as a pair.

    This helper deliberately does not create or modify ``mapping.json``.  The
    mapping is a structured artifact owned by the pipeline; this function only
    renders the two human-readable Markdown companions.
    """

    target_dir = Path(output_dir).expanduser()
    target_dir.mkdir(parents=True, exist_ok=True)
    draft_path = target_dir / "draft.docx"
    analysis_path = target_dir / "analysis.docx"
    render_markdown_file(draft_markdown, draft_path)
    render_markdown_file(analysis_markdown, analysis_path)
    return RenderedDocuments(draft_docx=draft_path, analysis_docx=analysis_path)


def _new_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)
    _configure_styles(document)
    document.core_properties.title = title
    document.core_properties.subject = "Source-grounded desktop procedure artifact"
    document.core_properties.author = "Document Enhancer"
    return document


def _configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER_PT)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING

    # Heading styles are real Word heading styles, not visually simulated
    # paragraphs.  This keeps the document navigable and preserves Markdown
    # section hierarchy for downstream consumers.
    heading_tokens = {
        1: (16.0, HEADING_COLOR, 18.0, 10.0),
        2: (13.0, HEADING_COLOR, 14.0, 7.0),
        3: (12.0, HEADING_DARK_COLOR, 10.0, 5.0),
        4: (11.5, HEADING_DARK_COLOR, 8.0, 4.0),
        5: (11.0, HEADING_DARK_COLOR, 7.0, 3.0),
        6: (11.0, HEADING_DARK_COLOR, 6.0, 3.0),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    title_style = styles["Title"]
    title_style.font.name = BODY_FONT
    title_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string("0B2545")
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(8)
    title_style.paragraph_format.line_spacing = 1.0
    title_style.paragraph_format.keep_with_next = True

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = BODY_FONT
    subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    subtitle_style.font.size = Pt(13.5)
    subtitle_style.font.italic = False
    subtitle_style.font.color.rgb = RGBColor.from_string("667085")
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(16)
    subtitle_style.paragraph_format.line_spacing = 1.0

    for list_style_name in ("List Paragraph", "List Bullet", "List Number"):
        list_style = styles[list_style_name]
        list_style.font.name = BODY_FONT
        list_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        list_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        list_style.font.size = Pt(BODY_SIZE_PT)
        list_style.font.color.rgb = RGBColor.from_string("1F2937")
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(4)
        list_style.paragraph_format.line_spacing = BODY_LINE_SPACING

    code_style = _get_or_add_style(document, "Code Block")
    code_style.font.name = "Courier New"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code_style.font.size = Pt(9.5)
    code_style.font.color.rgb = RGBColor.from_string("344054")
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.right_indent = Inches(0.2)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0

    table_style = _get_or_add_style(document, "Table Text")
    table_style.font.name = BODY_FONT
    table_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    table_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    table_style.font.size = Pt(9.5)
    table_style.font.color.rgb = RGBColor.from_string("1F2937")
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(0)
    table_style.paragraph_format.line_spacing = 1.0


def _get_or_add_style(document: Document, name: str):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_header_footer(document: Document, title: str) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    run = header_paragraph.add_run(title)
    _format_run(run, size=8.5, color=MUTED_COLOR, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    label = footer_paragraph.add_run("Document Enhancer  |  ")
    _format_run(label, size=8.0, color=MUTED_COLOR)
    _add_page_number_field(footer_paragraph)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    _format_run(run, size=8.0, color=MUTED_COLOR)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _render_blocks(document: Document, tokens: Sequence[Token], *, base_dir: Path | None) -> None:
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type.endswith("_close"):
            index += 1
            continue

        if token.type == "heading_open":
            inline = _next_inline(tokens, index)
            if inline is not None:
                level = _heading_level(token.tag)
                paragraph = document.add_paragraph(style=f"Heading {level}")
                _append_inline_tokens(paragraph, inline.children or (), base_dir=base_dir)
            index += 3
            continue

        if token.type == "paragraph_open":
            inline = _next_inline(tokens, index)
            if inline is not None and not token.hidden:
                paragraph = document.add_paragraph(style="Normal")
                children = inline.children or ()
                if _is_standalone_image(children):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _append_inline_tokens(paragraph, children, base_dir=base_dir)
            index += 3
            continue

        if token.type in {"bullet_list_open", "ordered_list_open"}:
            index = _render_list(document, tokens, index, level=0, base_dir=base_dir)
            continue

        if token.type == "table_open":
            index = _render_table(document, tokens, index, base_dir=base_dir)
            continue

        if token.type == "blockquote_open":
            index = _render_blockquote(document, tokens, index, base_dir=base_dir)
            continue

        if token.type in {"fence", "code_block"}:
            _render_code_block(document, token.content)
            index += 1
            continue

        if token.type == "hr":
            _render_horizontal_rule(document)
            index += 1
            continue

        # HTML blocks are not expected in generated artifacts.  Preserve any
        # meaningful text rather than dropping it silently.
        if token.type == "html_block" and token.content.strip():
            paragraph = document.add_paragraph(style="Normal")
            paragraph.add_run(token.content.strip())
            index += 1
            continue

        index += 1


def _heading_level(tag: str) -> int:
    try:
        level = int(tag.removeprefix("h"))
    except ValueError:
        return 1
    return max(1, min(6, level))


def _next_inline(tokens: Sequence[Token], index: int) -> Token | None:
    if index + 1 >= len(tokens):
        return None
    candidate = tokens[index + 1]
    return candidate if candidate.type == "inline" else None


def _render_list(
    document: Document,
    tokens: Sequence[Token],
    start: int,
    *,
    level: int,
    base_dir: Path | None,
) -> int:
    opening = tokens[start]
    ordered = opening.type == "ordered_list_open"
    closing_type = "ordered_list_close" if ordered else "bullet_list_close"
    list_start = _ordered_list_start(opening) if ordered else 1
    num_id = _new_numbering_instance(document, ordered=ordered, start=list_start)
    index = start + 1
    item_level: int | None = None
    first_paragraph = True

    while index < len(tokens):
        token = tokens[index]
        if token.type == closing_type and token.level == opening.level:
            return index + 1
        if token.type == "list_item_open":
            item_level = token.level
            first_paragraph = True
            index += 1
            continue
        if token.type == "list_item_close":
            item_level = None
            index += 1
            continue
        if token.type in {"bullet_list_open", "ordered_list_open"}:
            # A nested list belongs to the current item and gets its own
            # numbering instance, which correctly restarts ordered sublists.
            index = _render_list(document, tokens, index, level=level + 1, base_dir=base_dir)
            continue
        if token.type == "inline" and item_level is not None:
            paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
            _set_list_paragraph_geometry(paragraph, level=level, numbered=ordered)
            if first_paragraph:
                _apply_numbering(paragraph, num_id=num_id, level=level)
                first_paragraph = False
            else:
                # A loose list item's second paragraph remains aligned with its
                # text but deliberately has no second marker.
                paragraph.paragraph_format.left_indent = Inches(
                    LIST_TEXT_AT_IN + level * LIST_LEVEL_STEP_IN
                )
                paragraph.paragraph_format.first_line_indent = Inches(0)
            _append_inline_tokens(paragraph, token.children or (), base_dir=base_dir)
            index += 1
            continue
        index += 1
    raise RenderError("malformed Markdown list: missing closing list token")


def _set_list_paragraph_geometry(paragraph, *, level: int, numbered: bool) -> None:
    del numbered  # The numbering definition carries the marker's exact geometry.
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.left_indent = Inches(LIST_TEXT_AT_IN + level * LIST_LEVEL_STEP_IN)
    paragraph.paragraph_format.first_line_indent = Inches(-LIST_HANGING_IN)


def _apply_numbering(paragraph, *, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def _ordered_list_start(token: Token) -> int:
    attrs = token.attrs or {}
    raw_start = attrs.get("start", 1) if hasattr(attrs, "get") else 1
    try:
        return max(1, int(raw_start))
    except (TypeError, ValueError):
        return 1


def _new_numbering_instance(document: Document, *, ordered: bool, start: int = 1) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)
    for level in range(LIST_MAX_LEVEL):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start_element = OxmlElement("w:start")
        start_element.set(qn("w:val"), str(start))
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else "•")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(int((LIST_TEXT_AT_IN + level * LIST_LEVEL_STEP_IN) * 1440)))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(int((LIST_TEXT_AT_IN + level * LIST_LEVEL_STEP_IN) * 1440)))
        indent.set(qn("w:hanging"), str(int(LIST_HANGING_IN * 1440)))
        p_pr.extend([tabs, indent])
        lvl.extend([start_element, num_fmt, lvl_text, lvl_jc, p_pr])
        if not ordered:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), BODY_FONT)
            fonts.set(qn("w:hAnsi"), BODY_FONT)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _render_table(
    document: Document,
    tokens: Sequence[Token],
    start: int,
    *,
    base_dir: Path | None,
) -> int:
    rows: list[list[_InlineCell]] = []
    current_row: list[_InlineCell] | None = None
    cell_header = False
    index = start + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_close":
            break
        if token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close":
            if current_row:
                rows.append(current_row)
            current_row = None
        elif token.type in {"th_open", "td_open"}:
            cell_header = token.type == "th_open"
        elif token.type == "inline" and current_row is not None:
            current_row.append(
                _InlineCell(children=tuple(token.children or ()), header=cell_header)
            )
        index += 1
    if not rows:
        return index + 1

    column_count = max(len(row) for row in rows)
    for row in rows:
        row.extend(_InlineCell(children=(), header=False) for _ in range(column_count - len(row)))
    widths = _table_column_widths(rows, column_count)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_geometry(table, widths)
    _set_table_borders(table)
    _set_cell_margins(table)

    for row_index, row in enumerate(rows):
        for column_index, cell_data in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.width = Inches(widths[column_index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["Table Text"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_index == 0:
                _shade_cell(cell, TABLE_HEADER_FILL)
            _append_inline_tokens(paragraph, cell_data.children, table_cell=True, base_dir=base_dir)
            if cell_data.header:
                for run in paragraph.runs:
                    run.bold = True
    # Let the first row repeat if the table spans pages.
    _set_repeat_table_header(table.rows[0])
    for row in table.rows:
        _set_table_row_cant_split(row)
    return index + 1


def _table_column_widths(rows: Sequence[Sequence[_InlineCell]], count: int) -> list[int]:
    lengths = [1] * count
    for row in rows:
        for index, cell in enumerate(row):
            text = "".join(
                token.content for token in cell.children if token.type in {"text", "code_inline"}
            )
            lengths[index] = max(lengths[index], min(64, len(text)))
    minimum = 1080
    available = CONTENT_WIDTH_DXA - minimum * count
    total = sum(lengths)
    widths = [minimum + (available * length // total) for length in lengths]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_table_geometry(table, widths: Sequence[int]) -> None:
    table_element = table._tbl
    tbl_pr = table_element.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(sum(widths)))

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width))


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER_COLOR)


def _set_cell_margins(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (
                ("top", CELL_MARGIN_TOP_DXA),
                ("bottom", CELL_MARGIN_BOTTOM_DXA),
                ("start", CELL_MARGIN_START_DXA),
                ("end", CELL_MARGIN_END_DXA),
            ):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _render_blockquote(
    document: Document,
    tokens: Sequence[Token],
    start: int,
    *,
    base_dir: Path | None,
) -> int:
    # Blockquote content is rendered in document order with a restrained left
    # rule.  Paragraphs remain ordinary paragraphs, so extraction does not lose
    # their text or inline emphasis.
    opening = tokens[start]
    index = start + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == "blockquote_close" and token.level == opening.level:
            return index + 1
        if token.type == "paragraph_open":
            inline = _next_inline(tokens, index)
            if inline is not None:
                paragraph = document.add_paragraph(style="Normal")
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.right_indent = Inches(0.1)
                _set_paragraph_left_border(paragraph)
                _append_inline_tokens(paragraph, inline.children or (), base_dir=base_dir)
            index += 3
            continue
        if token.type in {"bullet_list_open", "ordered_list_open"}:
            index = _render_list(document, tokens, index, level=1, base_dir=base_dir)
            continue
        index += 1
    raise RenderError("malformed Markdown blockquote: missing closing token")


def _render_code_block(document: Document, content: str) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    _shade_paragraph(paragraph, CODE_FILL)
    lines = content.rstrip("\n").splitlines() or [""]
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        _format_run(run, name="Courier New", size=9.5, color="344054")
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def _render_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), TABLE_BORDER_COLOR)
    borders.append(bottom)
    p_pr.append(borders)


def _set_paragraph_left_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), HEADING_COLOR)
    borders.append(left)
    p_pr.append(borders)


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _append_inline_tokens(
    paragraph,
    children: Iterable[Token],
    *,
    table_cell: bool = False,
    base_dir: Path | None,
) -> None:
    bold = False
    italic = False
    strike = False
    link = False
    for token in children:
        token_type = token.type
        if token_type == "strong_open":
            bold = True
            continue
        if token_type == "strong_close":
            bold = False
            continue
        if token_type == "em_open":
            italic = True
            continue
        if token_type == "em_close":
            italic = False
            continue
        if token_type in {"s_open", "del_open"}:
            strike = True
            continue
        if token_type in {"s_close", "del_close"}:
            strike = False
            continue
        if token_type == "link_open":
            link = True
            continue
        if token_type == "link_close":
            link = False
            continue
        if token_type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break(WD_BREAK.LINE)
            continue
        if token_type == "code_inline":
            run = paragraph.add_run(token.content)
            _format_run(
                run,
                name="Courier New",
                size=9.5 if table_cell else 10.0,
                color="344054",
            )
            _apply_inline_state(run, bold=bold, italic=italic, strike=strike, link=link)
            _shade_run(run, CODE_FILL)
            continue
        if token_type == "image":
            _append_image(paragraph, token, base_dir=base_dir)
            continue
        if token_type in {"html_inline", "text"}:
            content = token.content
            if content:
                run = paragraph.add_run(content)
                _apply_inline_state(run, bold=bold, italic=italic, strike=strike, link=link)
            continue
        if token_type == "entity":
            run = paragraph.add_run(token.content)
            _apply_inline_state(run, bold=bold, italic=italic, strike=strike, link=link)
            continue
        # ``emoji``, inline math, or another plugin token: retain its literal
        # content when present rather than dropping source information.
        if token.content:
            run = paragraph.add_run(token.content)
            _apply_inline_state(run, bold=bold, italic=italic, strike=strike, link=link)


def _is_standalone_image(children: Sequence[Token]) -> bool:
    return len(children) == 1 and children[0].type == "image"


def _append_image(paragraph, token: Token, *, base_dir: Path | None) -> None:
    attrs = token.attrs or {}
    source = attrs.get("src", "") if hasattr(attrs, "get") else ""
    image_path = _resolve_image_path(source, base_dir=base_dir)
    try:
        shape = paragraph.add_run().add_picture(str(image_path))
    except Exception as exc:  # image decoders expose several format-specific errors
        raise RenderError(
            f"unsupported or unreadable Markdown image {source!r} at {image_path}: {exc}"
        ) from exc

    max_width = int(Inches(CONTENT_WIDTH_DXA / 1440))
    shape_width = int(shape.width)
    if shape_width > max_width and shape_width > 0:
        shape.height = max(1, int(int(shape.height) * max_width / shape_width))
        shape.width = max_width

    alt = token.content or (attrs.get("alt", "") if hasattr(attrs, "get") else "")
    if alt:
        # ``descr`` is the OOXML alternate-description field exposed by Word's
        # accessibility tools.  Keep the title too because some readers use it
        # as the image tooltip instead of the description.
        shape._inline.docPr.set("descr", alt)
        shape._inline.docPr.set("title", alt)


def _resolve_image_path(source: str, *, base_dir: Path | None) -> Path:
    if not source:
        raise RenderError("Markdown image has no local source path")

    parsed = urlparse(source)
    if parsed.scheme or parsed.netloc or source.startswith("//"):
        raise RenderError(f"remote image URLs are not supported: {source}")
    if base_dir is None:
        raise RenderError(
            f"cannot resolve local Markdown image {source!r}: no Markdown source directory"
        )

    image_path = Path(unquote(source))
    if not image_path.is_absolute():
        image_path = base_dir / image_path
    try:
        available = image_path.is_file()
    except OSError as exc:
        raise RenderError(f"could not access Markdown image {image_path}: {exc}") from exc
    if not available:
        raise RenderError(f"Markdown image does not exist: {image_path}")
    return image_path


def _apply_inline_state(run, *, bold: bool, italic: bool, strike: bool, link: bool) -> None:
    """Apply only inline semantics, leaving the paragraph style authoritative.

    Heading runs must inherit their heading size/color/weight from the real
    Word heading style.  Writing ``bold=False`` or an 11 pt body color directly
    onto those runs would silently erase that style, so this helper writes only
    active inline attributes.
    """

    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if strike:
        run.font.strike = True
    if link:
        run.font.color.rgb = RGBColor.from_string(HEADING_COLOR)
        run.underline = True


def _format_run(
    run,
    *,
    name: str = BODY_FONT,
    size: float = BODY_SIZE_PT,
    color: str = "1F2937",
    bold: bool | None = None,
    italic: bool | None = None,
    strike: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if strike is not None:
        run.font.strike = strike
    if underline is not None:
        run.underline = underline


def _shade_run(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    r_pr.append(shading)


def _title_from_markdown(markdown: str) -> str:
    """Extract the first level-one heading for metadata and page furniture."""

    tokens = _parser().parse(markdown)
    for index, token in enumerate(tokens[:-1]):
        if token.type == "heading_open" and token.tag == "h1":
            inline = tokens[index + 1]
            if inline.type == "inline":
                return inline.content.strip()
    return ""
